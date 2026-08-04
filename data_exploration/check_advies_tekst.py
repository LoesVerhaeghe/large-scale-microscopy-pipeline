from pathlib import Path
import shutil
from docx import Document



output_folder = Path("QIA/experiments_export")
output_folder.mkdir(parents=True, exist_ok=True)


for order_nr, group in df.groupby("order_nr"):

    # create folder for this experiment
    experiment_folder = output_folder / str(order_nr)
    experiment_folder.mkdir(exist_ok=True)

    print(f"Processing order {order_nr}")

    # -------------------------
    # Save advies tekst to Word
    # -------------------------

    doc = Document()

    # if multiple rows have text, combine them
    texts = group["advies_tekst"].dropna().unique()

    for text in texts:
        doc.add_paragraph(str(text))
        doc.add_paragraph("")  # empty line between texts

    doc_path = experiment_folder / "advies_tekst.docx"
    doc.save(doc_path)


    # -------------------------
    # Copy images
    # -------------------------

    for _, row in group.iterrows():

        image_path = Path(row["image_path"])

        if image_path.exists():

            destination = experiment_folder / image_path.name

            shutil.copy2(
                image_path,
                destination
            )

        else:
            print(f"Image not found: {image_path}")

print("Done!")